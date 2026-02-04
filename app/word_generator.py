"""
Module to generate Word documents for reports in Vietnamese administrative format.
"""
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.models import Task, TaskStatus
from app.config import config

logger = logging.getLogger(__name__)


class WordReportGenerator:
    """Generate Word documents for various report types."""
    
    def __init__(self, output_dir: str = "reports"):
        """
        Initialize the Word report generator.
        
        Args:
            output_dir: Directory to save generated reports
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        logger.info(f"Word generator initialized. Output directory: {self.output_dir}")
    
    def _setup_document_styles(self, doc: Document):
        """Setup Vietnamese administrative document styles."""
        # Title style
        styles = doc.styles
        
        # Header style for organization name
        if 'Header Style' not in styles:
            header_style = styles.add_style('Header Style', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Times New Roman'
            header_style.font.size = Pt(11)
            header_style.font.bold = True
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title style
        if 'Report Title' not in styles:
            title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_before = Pt(12)
            title_style.paragraph_format.space_after = Pt(12)
    
    def _add_header(self, doc: Document, title: str):
        """Add Vietnamese administrative document header."""
        # Organization header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("VIỆN CÔNG NGHỆ SỐ\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        
        run2 = header.add_run("TỔ THƯ KÝ")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
        run2.font.bold = True
        
        # Horizontal line
        doc.add_paragraph("_" * 35).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.upper())
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Ngày {datetime.now().strftime('%d tháng %m năm %Y')}")
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(11)
        date_run.italic = True
        
        doc.add_paragraph()  # Empty line
    
    def _add_footer(self, doc: Document):
        """Add signature section to document."""
        doc.add_paragraph()  # Empty line
        
        # Signature section
        signature_para = doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = signature_para.add_run("PHỤ TRÁCH TỔ THƯ KÝ\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        
        # Space for signature
        for _ in range(3):
            empty = doc.add_paragraph()
            empty.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Name placeholder
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        name_run = name_para.add_run("(Ký và ghi rõ họ tên)")
        name_run.font.name = 'Times New Roman'
        name_run.font.size = Pt(11)
        name_run.italic = True
    
    def _add_task_table(self, doc: Document, tasks: List[Task]):
        """Add a formatted table of tasks to the document."""
        if not tasks:
            para = doc.add_paragraph("Không có công việc nào.")
            para.runs[0].font.name = 'Times New Roman'
            para.runs[0].font.size = Pt(11)
            para.runs[0].italic = True
            return
        
        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['STT', 'Họ tên', 'Nội dung công việc', 'Mức độ', 'Deadline', 'Tiến độ']
        
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            # Format header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for task in tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = str(task.stt)
            row_cells[1].text = task.ho_ten or ''
            row_cells[2].text = task.noi_dung or ''
            row_cells[3].text = task.muc_do or ''
            row_cells[4].text = task.deadline.strftime('%d/%m/%Y') if task.deadline else ''
            row_cells[5].text = task.ket_qua or ''
            
            # Format data cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
        
        # Adjust column widths
        table.columns[0].width = Inches(0.5)  # STT
        table.columns[1].width = Inches(1.5)  # Họ tên
        table.columns[2].width = Inches(2.5)  # Nội dung
        table.columns[3].width = Inches(1.0)  # Mức độ
        table.columns[4].width = Inches(1.0)  # Deadline
        table.columns[5].width = Inches(1.5)  # Tiến độ
    
    def _add_grouped_tasks(self, doc: Document, grouped_tasks: Dict[TaskStatus, List[Task]]):
        """Add grouped tasks by status to document."""
        status_names = {
            TaskStatus.OVERDUE: "Công việc quá hạn",
            TaskStatus.DUE_TODAY: "Công việc đến hạn hôm nay",
            TaskStatus.DUE_TOMORROW: "Công việc đến hạn ngày mai",
            TaskStatus.DUE_2_3_DAYS: "Công việc đến hạn trong 2-3 ngày",
            TaskStatus.ON_TRACK: "Công việc đang theo kế hoạch",
            TaskStatus.NO_DEADLINE: "Công việc không có deadline"
        }
        
        for status in [TaskStatus.OVERDUE, TaskStatus.DUE_TODAY, TaskStatus.DUE_TOMORROW, 
                       TaskStatus.DUE_2_3_DAYS, TaskStatus.ON_TRACK]:
            if status in grouped_tasks and grouped_tasks[status]:
                # Section header
                section_para = doc.add_paragraph()
                section_run = section_para.add_run(f"\n{status_names[status]} ({len(grouped_tasks[status])} việc)")
                section_run.font.name = 'Times New Roman'
                section_run.font.size = Pt(13)
                section_run.font.bold = True
                
                # Add tasks table
                self._add_task_table(doc, grouped_tasks[status])
    
    def generate_daily_report(self, tasks: List[Task], grouped_tasks: Dict[TaskStatus, List[Task]]) -> Path:
        """
        Generate daily report as Word document.
        
        Args:
            tasks: All incomplete tasks
            grouped_tasks: Tasks grouped by status
            
        Returns:
            Path to the generated Word document
        """
        doc = Document()
        self._setup_document_styles(doc)
        
        # Add header
        self._add_header(doc, "Báo cáo tiến độ công việc hàng ngày")
        
        # Summary
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(f"Tổng số công việc: {len(tasks)}\n")
        summary_run.font.name = 'Times New Roman'
        summary_run.font.size = Pt(11)
        
        overdue_count = len(grouped_tasks.get(TaskStatus.OVERDUE, []))
        if overdue_count > 0:
            overdue_run = summary_para.add_run(f"⚠️ Cảnh báo: Có {overdue_count} công việc quá hạn\n")
            overdue_run.font.name = 'Times New Roman'
            overdue_run.font.size = Pt(11)
            overdue_run.font.color.rgb = RGBColor(255, 0, 0)
            overdue_run.font.bold = True
        
        doc.add_paragraph()
        
        # Add grouped tasks
        self._add_grouped_tasks(doc, grouped_tasks)
        
        # Add footer
        self._add_footer(doc)
        
        # Save document
        filename = f"bao_cao_ngay_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        logger.info(f"Daily report generated: {filepath}")
        return filepath
    
    def generate_weekly_report(self, tasks: List[Task], grouped_tasks: Dict[TaskStatus, List[Task]]) -> Path:
        """
        Generate weekly report as Word document.
        
        Args:
            tasks: All incomplete tasks
            grouped_tasks: Tasks grouped by status
            
        Returns:
            Path to the generated Word document
        """
        doc = Document()
        self._setup_document_styles(doc)
        
        # Add header
        self._add_header(doc, "Báo cáo tiến độ công việc tuần")
        
        # Week info
        now = datetime.now()
        week_para = doc.add_paragraph()
        week_run = week_para.add_run(f"Tuần {now.isocalendar()[1]} - Năm {now.year}\n")
        week_run.font.name = 'Times New Roman'
        week_run.font.size = Pt(11)
        week_run.font.bold = True
        
        # Summary
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(f"Tổng số công việc: {len(tasks)}\n")
        summary_run.font.name = 'Times New Roman'
        summary_run.font.size = Pt(11)
        
        overdue_count = len(grouped_tasks.get(TaskStatus.OVERDUE, []))
        due_soon = len(grouped_tasks.get(TaskStatus.DUE_TODAY, [])) + \
                   len(grouped_tasks.get(TaskStatus.DUE_TOMORROW, [])) + \
                   len(grouped_tasks.get(TaskStatus.DUE_2_3_DAYS, []))
        
        if overdue_count > 0:
            overdue_run = summary_para.add_run(f"⚠️ Quá hạn: {overdue_count} công việc\n")
            overdue_run.font.name = 'Times New Roman'
            overdue_run.font.size = Pt(11)
            overdue_run.font.color.rgb = RGBColor(255, 0, 0)
        
        if due_soon > 0:
            due_run = summary_para.add_run(f"⏰ Sắp đến hạn: {due_soon} công việc\n")
            due_run.font.name = 'Times New Roman'
            due_run.font.size = Pt(11)
            due_run.font.color.rgb = RGBColor(255, 165, 0)
        
        doc.add_paragraph()
        
        # Add grouped tasks
        self._add_grouped_tasks(doc, grouped_tasks)
        
        # Add footer
        self._add_footer(doc)
        
        # Save document
        filename = f"bao_cao_tuan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        logger.info(f"Weekly report generated: {filepath}")
        return filepath
    
    def generate_overdue_report(self, overdue_by_person: Dict[str, List[Task]]) -> Path:
        """
        Generate overdue tasks report by person.
        
        Args:
            overdue_by_person: Dictionary mapping person name to their overdue tasks
            
        Returns:
            Path to the generated Word document
        """
        doc = Document()
        self._setup_document_styles(doc)
        
        # Add header
        self._add_header(doc, "Báo cáo công việc quá hạn")
        
        # Summary
        total_overdue = sum(len(tasks) for tasks in overdue_by_person.values())
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(f"⚠️ Tổng số công việc quá hạn: {total_overdue}\n")
        summary_run.font.name = 'Times New Roman'
        summary_run.font.size = Pt(11)
        summary_run.font.bold = True
        summary_run.font.color.rgb = RGBColor(255, 0, 0)
        
        people_run = summary_para.add_run(f"Số người có công việc quá hạn: {len(overdue_by_person)}\n")
        people_run.font.name = 'Times New Roman'
        people_run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Add tasks by person
        if not overdue_by_person:
            para = doc.add_paragraph("✅ Không có công việc quá hạn!")
            para.runs[0].font.name = 'Times New Roman'
            para.runs[0].font.size = Pt(11)
            para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        else:
            for person, tasks in sorted(overdue_by_person.items()):
                # Person header
                person_para = doc.add_paragraph()
                person_run = person_para.add_run(f"\n{person} - {len(tasks)} công việc quá hạn")
                person_run.font.name = 'Times New Roman'
                person_run.font.size = Pt(12)
                person_run.font.bold = True
                
                # Tasks table
                self._add_task_table(doc, tasks)
        
        # Add footer
        self._add_footer(doc)
        
        # Save document
        filename = f"bao_cao_qua_han_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        logger.info(f"Overdue report generated: {filepath}")
        return filepath
